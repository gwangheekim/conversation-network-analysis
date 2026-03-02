import os
import pandas as pd
import docx
import re
from tqdm import tqdm
import time
import argparse
import json
# External LLM providers
import openai
from openai import OpenAI
import anthropic
# import google.generativeai as genai
from google import genai
from google.genai import types

"""
python llm_api.py --provider "[provider (google, openai, anthropic)]" --apikey_json "[apikey path]" --model_name "[model name]" --transcript_dir "[data path]" --output_csv "[output file name]"
"""

# --- 1. Docx parser ---
def extract_dialogue_from_docx(docx_path):
    doc = docx.Document(docx_path)
    full_text = [paragraph.text for paragraph in doc.paragraphs]
    text = '\n'.join(full_text)

    chunks = re.split(
        r'(\n+\d{1,2}:\d{2}(?::\d{2})?\s*\n|\n+\d{1,2}:\d{2}(?::\d{2})?\s*$|^\d{1,2}:\d{2}(?::\d{2})?\s*\n)',
        text, flags=re.MULTILINE
    )

    timestamps, speakers, utterances = [], [], []
    current_timestamp = None

    for chunk in chunks:
        clean_chunk = chunk.strip('\n').strip()
        if re.fullmatch(r"\d{1,2}:\d{2}(?::\d{2})?", clean_chunk):
            current_timestamp = clean_chunk
        else:
            lines = [l.strip() for l in chunk.split('\n') if l.strip()]
            i = 0
            while i < len(lines):
                line = lines[i]
                match = re.match(r'^([A-Za-z가-힣]+):\s*(.+)', line)
                if match:
                    speaker, utterance = match.group(1), match.group(2)
                    utterance = re.sub(r'\s+', ' ', utterance)
                    timestamps.append(current_timestamp)
                    speakers.append(speaker)
                    utterances.append(utterance)
                elif re.match(r'^\d{1,2}:\d{2}(?::\d{2})?\s+\[.+\]$', line):
                    ts, desc = re.match(r'^(\d{1,2}:\d{2}(?::\d{2})?)\s+(\[.+\])$', line).groups()
                    current_timestamp = ts
                    timestamps.append(current_timestamp)
                    speakers.append("NARRATION")
                    utterances.append(desc)
                elif re.match(r'^\[.+\]$', line):
                    timestamps.append(current_timestamp)
                    speakers.append("NARRATION")
                    utterances.append(line)
                i += 1
    df = pd.DataFrame({'timestamp': timestamps, 'name': speakers, 'utterance': utterances})
    df = df[df.name != "NARRATION"].reset_index(drop=True)
    return df


# --- 2. Prompt ---
def make_prompt_messages(dialogue_info, utterance_id_for_output):
    system_prompt = f"""
    You are an expert in analyzing utterance contributions in problem-solving discussions. Your task is to perform two steps:

    IMPORTANT NOTE ABOUT UTTERANCE CONTENT:
    - Content within parentheses () or square brackets [] represents ACTIONS/BEHAVIORS, not speech.
    - Example: "[points to the board]", "(writes on paper)", "[moves blocks]" are actions, not verbal utterances.
    - If an utterance consists entirely of action descriptions with no speech, classify it as Uncorrelated (UC).

    STEP 1: REASONING
    Analyze the utterance based on the rules below. BRIEFLY explain your reasoning (1-2 concise sentences) on its contribution to idea development.
    Example: Reasoning: Speaker explains their method of adding 10s, fitting EXP criteria as it details their reasoning.

    STEP 2: CLASSIFICATION OUTPUT
    Provide the classification as a SINGLE line: "FINAL_CLASSIFICATION: [UTTERANCE_ID] : [NAME] : [LABEL] : [REFERENCE]"
    - EXP & UC: [REFERENCE] is 'NA'.
    - EOI: [REFERENCE] is the NAME of the participant whose idea is engaged (trace back if needed). Use 'Previous Speaker' if name is unknown but clearly refers to t-1 idea.

    --- DECISION FLOW ---
    Follow this order strictly:
    1.  Check for EOI (Engage Others Idea) first. Does the utterance DIRECTLY respond to, critique, correct, or question a specific, conceptual idea from a peer? If YES, classify as EOI and determine the level (LOW, MEDIUM, HIGH).
    2.  If not EOI, check for EXP (Explain Own Idea). Is the speaker verbalizing their own conceptual reasoning, strategy, or the 'why' behind their actions? If YES, classify as EXP.
    3.  If not EOI or EXP, classify as UC (Uncorrelated). This is the default for anything that lacks clear conceptual contribution.

    --- CRITICAL WARNING: AVOID COMMON MISTAKES ---
    - UC vs. EOI(LOW): A simple "Yes", "No", "I agree", or "Wow!" is NOT EOI. It is Uncorrelated (UC) unless it clearly rephrases the peer's conceptual idea, proving understanding. Social agreement alone is not conceptual engagement.
    - UC vs. EXP: Simply describing an action ("I'm counting") or stating an answer ("It's 35") is Uncorrelated (UC). To be EXP, the speaker must explain the purpose of the action ("I'm counting to check my answer") or the reasoning behind the answer ("It's 35 because 70 divided by 2 is 35").

    --- CATEGORY DEFINITIONS & CRITERIA ---

    1.  Explain Own Idea (EXP):
        -   Definition: Speaker verbalizes their own conceptual reasoning, mathematical strategy, or understanding with clear explanation, justification, or elaboration that reveals their thought process. This includes explanations of the speaker's own ideas or actions, even if prompted by a peer's question.
        -   Key Includes:
            *   Elaborating a novel strategy or a significant modification to an existing one.
            *   Explaining a conceptual realization with its basis.
            *   Justifying a mathematical method or answer.
            *   Articulating a step-by-step conceptual plan or calculation process with its underlying logic.

    2.  Engage Others Idea (EOI): Speaker directly engages with the content of another's specific articulated conceptual mathematical idea. The level depends on the depth and nature of this engagement.

        -   EOI(LOW): Minimum Conceptual Engagement - Confirmation & Basic Clarification
            *   Definition: Speaker briefly acknowledges, confirms, agrees/disagrees with, or asks a very simple, direct question about a clearly identifiable, specific conceptual mathematical point previously articulated by a peer. The utterance MUST demonstrate BOTH (1) clear understanding of that specific conceptual point AND (2) make a direct, unambiguous verbal reference to it.
            *   Key Indicators: Confirmatory Restatement, Targeted Agreement/Disagree-ment, Simple clarification on a specific detail within the peer's idea.

        -   EOI(MEDIUM): Probing & Minor Elaboration (of a peer's idea)
            *   Definition: Speaker seeks to deepen their understanding of a peer's specific conceptual idea by asking questions that probe the underlying reasoning, assumptions, or implications of that peer's idea. They might also restate the peer's idea with minor conceptual additions or offer small, scaffolded extensions.
            *   Key Indicators: Probing "Why" or "How" Questions about a peer's reasoning, Requesting Elaboration on a peer's idea, Identifying Ambiguity for clarification within a peer's idea.

        -   EOI(HIGH): Critique, Significant Extension, & Synthesis (of a peer's idea)
            *   Definition: Speaker critically evaluates a peer's specific conceptual idea by offering a reasoned critique, proposing a significant and justified alternative/refinement to that peer's idea, or by synthesizing it with other ideas to create a more robust solution. This involves substantial conceptual work.
            *   Clarification on Critique: A critique must be reasoned. A simple, unelaborated disagreement or correction (e.g., 'No, the answer is 20') does not qualify as EOI(HIGH). It is a targeted disagreement about a conceptual point but lacks justification, making it EOI(LOW). To be EOI(HIGH), the utterance must explain why the peer's idea is incorrect (e.g., 'No, it's 20, because you forgot to add the 2.'), or clearly justify their disagreement by demonstrating their own correct reasoning or solution process (verbally or through action) in a manner that implicitly highlights the peer's error in relation to their original idea.
            *   Key Indicators: Reasoned Conceptual Critique of a peer's idea, Correction with Justification related to a peer's error, Proposing Justified Alternatives to a peer's idea, Conceptual Synthesis of a peer's idea with others.
            *   Proposing a Conceptual Paradigm Shift: This applies when a speaker clearly rejects a peer's more concrete, lower-level strategy and proposes a shift to a much more abstract or higher-level conceptual category. In this case, even if the proposed alternative does not include a specific methodology, the 'leap' itself is considered a significant critique and proposal that fundamentally changes the direction of problem-solving.

    3.  Uncorrelated (UC):
        -   Definition: No meaningful or discernible conceptual mathematical contribution; utterance is vague, off-task, purely procedural without explanation, or a brief unelaborated realization/intention.
        -   Key Includes:
            *   Narration of one's own actions without explaining the conceptual purpose.
            *   Vague statements ("This is hard," "I'm confused" without specifying the conceptual point).
            *   Simple affirmations/negations ("Yes", "I agree", "Okay") without specific reference to a peer's conceptual idea and demonstrating understanding of it.
            *   Generic procedural questions/statements ("What's next?").
            *   Unelaborated realizations ("I get it!", "I know!", "Oh!").
            *   Stating answers or numbers without explaining the process or reasoning.
            *   Reading aloud from a worksheet or simply repeating a peer's words without adding conceptual value.
            *   Brief self-corrections ('Oops') without explaining the conceptual reason for the correction.

    --- CORE CLASSIFICATION PRINCIPLES ---
    1. Conceptual Contribution Focus: Your primary focus is on conceptual idea development in mathematics.
    2. EOI - Engaging a Peer's Conceptual Idea: EOI requires engagement with a specific peer's conceptual mathematical idea, not just their words or a surface-level answer.
    3. EXP - Explaining Own Conceptual Reasoning: EXP requires explaining one's own reasoning, strategy, or understanding with clear justification or elaboration.
    4. Priority for Primary Conceptual Function: Choose the category that best reflects the utterance's dominant conceptual purpose. In complex cases where functions are mixed, use the Tie-Breaker rules below to help determine this primary function.
    5. UC as Default, with Caution for EOI(LOW): If an utterance does not clearly meet the criteria for EXP or EOI, classify it as UC. However, do not default to UC too aggressively if it meets the minimum criteria for EOI(LOW).
    6. Context is Key: Your classification must be based on the flow of the dialogue. Look beyond the immediate previous utterance; trace back to the origin of an idea for accurate EOI referencing or to understand the full context for an EXP.
    7. The "Target of Reasoning" Principle (EOI Precedence Rule): If an utterance's primary conceptual function is to respond directly to, critique, correct, or build upon a specific peer's articulated idea, it MUST be classified as EOI.
        - This is a strict priority rule. It holds true even if the speaker explains their own reasoning or calculation process as the means for that critique or extension. The target of the idea (a peer) takes precedence over the form of the reasoning (one's own calculation).
        - Critical Example (Correction as EOI(HIGH)):
            Peer A: "I got 18."
            Speaker B: "No, 9 plus 9 is 18, but you forgot to add the 2, so it's 20."
            Correct Classification: EOI(HIGH) : Peer A. This is NOT EXP. Although Speaker B explains their own calculation, the entire purpose of the utterance is to correct Peer A's idea.
    8. Principle of Conceptual Impact Over Explicit Explanation:
        - Do not automatically classify an utterance as UC just because it lacks an explicit "because" clause. Evaluate its function within the dialogue.
        - The Test: Does the utterance introduce a new, relevant mathematical object, operation, or logical step into the conversation? If yes, it likely has conceptual value.
            - Lean towards UC if it's just stating a number without context. (e.g., Peer A: "This is hard." Speaker B: "70.")
            - Lean towards EXP or EOI if it introduces a key operation or a logical next step based on the current context. (e.g., After the group agrees the total is 70 -> Speaker B: "Okay, so 70 divided by 2.") This utterance, despite its brevity, introduces the crucial next conceptual step (division) and is therefore not UC.
    9. Principle of Utterance Continuity: If an utterance is a direct continuation of the speaker's own previous turn (e.g., finishing a sentence or elaborating on a point immediately after pausing), evaluate the current utterance in the context of the combined statement. The classification should reflect the conceptual contribution of the complete idea, not just the final fragment.

    --- TIE-BREAKER RULES FOR BORDERLINE CASES ---
    The following are priority principles to increase consistency when classifying ambiguous, borderline cases.

    1. (EOI vs. EXP) The "Function over Form" Rule (Highest Priority):
    - This rule takes precedence over all others when an utterance has mixed characteristics. It applies when an utterance contains the speaker's own reasoning or calculation (form of EXP) but is used to directly critique, correct, or offer a counter-proposal to a peer's specific idea (function of EOI).
    - The Test: Ask, "What is the primary purpose of this utterance in the dialogue?"
        - If the purpose is to correct a peer's error, challenge a peer's method, or propose an alternative in direct response to a peer, it MUST be classified as EOI(HIGH). The corrective/critical function outweighs the explanatory form.
    - Critical Failure Case to Avoid:
        - Peer A: "I got the wrong answer, 18."
        - Speaker B: "No, because 9 plus 9 is 18, but you forgot to add the 2, so the answer is 20."
        - Correct Classification: Engage Others Idea (HIGH) : Peer A.
        - Incorrect Classification: Explain Own Idea : NA. Do not classify this as EXP simply because Speaker B explains their own calculation. The entire utterance exists only because of Peer A's mistake and its primary function is to correct it.

    2. (UC vs. EOI(LOW)) The "Rephrasing vs. Repeating" Rule:
    - An utterance that simply repeats a key word (e.g., "multiplication") or number from a peer should lean towards UC. This is because it provides insufficient evidence of conceptual understanding.
    - Conversely, classify as EOI(LOW) if the speaker rephrases the peer's idea in their own words (even slightly) or includes a specific result based on that idea (e.g., "Oh, so you mean 3 times 4 is 12"). This proves the speaker actively processed the idea.
    - The Test: Did the speaker parrot the peer's idea, or did they digest and explain it in their own head? Only the latter qualifies as EOI(LOW).

    3. (UC vs. EXP) The "Describing Actions vs. Explaining Purpose" Rule:
    - Simply describing an action ("I'm grouping these by 10s") is UC.
    - Explaining the conceptual purpose or reason (the 'Why') behind the action makes it EXP (e.g., "I'm grouping these by 10s to make them easier to count later").
    - Clarification on "Purpose": Explaining how a mathematical strategy is being constructed or how one is modeling the problem's components is also considered "explaining the purpose." The "why" is inherent in the strategic choice of representation.
        - Example of Conceptual Modeling (EXP): "So I'm putting 70 blocks here, and I'm making two groups, to show the total for the two days."
        - This is EXP because it explains how the actions (making groups, using 70 blocks) map onto the concepts of the problem (two days, total).

    4. (EXP vs. EOI(HIGH)) The "Foundation vs. Spark" Rule (for judging idea ownership):
    - To decide, ask this question: "Without the peer's idea, could the speaker's idea still stand on its own?"
    - EOI(HIGH) (Foundation Principle): This applies when the peer's idea serves as the 'foundation' for the speaker's idea. The speaker is modifying, critiquing, or extending the peer's idea, and without it, the speaker's utterance would lose its meaning.
    - EXP (Spark Principle): This applies when the peer's idea only served as a 'spark' for the speaker to come up with a conceptually independent idea. The speaker's idea would still stand as a complete strategy on its own.

    5. (EOI(LOW) vs. EOI(MEDIUM)) The "New vs. Implied Information" Test:
    - Lean towards EOI(LOW): if the utterance primarily confirms or rephrases information that was already stated or clearly implied by the peer. It demonstrates understanding of the existing idea. (e.g., Peer: "We need to split 70 in half." Speaker: "Oh, so 70 divided by 2.")
    - Lean towards EOI(MEDIUM): if the utterance adds a new, explicit piece of conceptual information (like a label, context, or reason) that was not present in the peer's original statement. It enriches the existing idea. (e.g., Peer: "We need to use the 2." Speaker: "Right, because the 2 represents the days.")
    """

    user_prompt = f"""
    Your task:
    Using the rules from the system prompt, classify the utterance with ID {utterance_id_for_output}. First give a brief reasoning, then provide a single line classification in the required format.

    Your response must include:
    1. A concise reasoning (STEP 1)
    2. One strict "FINAL_CLASSIFICATION:" line (STEP 2)
    
    Output format examples:
    Reasoning: Speaker explains own detailed method for grouping and multiplying.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Julian : Explain Own Idea : NA

    Reasoning: Speaker asks "Why did you decide to multiply there?" about John's specific step in his solution.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Julian : Engage Others Idea (MEDIUM) : John

    Reasoning: Speaker is just counting aloud "1, 2, 3" while moving blocks, no explanation of the conceptual plan.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Delilah : Uncorrelated : NA

    Reasoning: Peer A: "I think the pattern is adding 3 each time." Speaker B: "Yes, I see that too, it's adding 3." This shows understanding of and agreement with Peer A's specific conceptual pattern.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Keila : Engage Others Idea (LOW) : Peer A

    Reasoning: Peer A: "The answer is 12." Speaker B: "I agree." This is a generic agreement without referencing the conceptual method or reasoning of Peer A.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Keila : Uncorrelated : NA   

    Reasoning: Peer A states, "The result is 45." Speaker B says, "No, it's 50." This is a direct refutation of a peer's conceptual point (the result) but provides no reasoning or justification, qualifying it as low-level engagement.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Kenji : Engage Others Idea (LOW) : Peer A

    Reasoning: Speaker's utterance "...so we divide by two" completes their thought from their previous utterance "To find the amount for one day...". The combined idea explains the strategy of dividing by two to find the daily amount, making it EXP.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Sarah : Explain Own Idea : NA

    Reasoning: Speaker states "I know how to solve this now!" but doesn't explain the method or what they realized conceptually.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Alex : Uncorrelated : NA

    Reasoning: Speaker critiques Chris's method: "Your addition is correct, but multiplying 5 by 3 would be a more direct conceptual approach here."
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Ben : Engage Others Idea (HIGH) : Chris

    Reasoning: Speaker says "Okay, let's start problem 2." This is procedural coordination.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Maria : Uncorrelated : NA

    Reasoning: Speaker asks a standalone "What?" after a peer's long explanation, without indicating what part is unclear.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : David : Uncorrelated : NA

    Reasoning: Peer C: "So we need to find the total area." Speaker D: "So, by 'total area', you mean we should multiply length by width for each rectangle?" This clarifies understanding of the peer's specific conceptual goal.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Chloe : Engage Others Idea (LOW) : Peer C

    Reasoning: Speaker explains their realization: "Ah, now I understand! 'Each' means we need to perform the operation for every single item, not just once for the total." This explains what was understood conceptually.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Sam : Explain Own Idea : NA

    Reasoning: Speaker says "My strategy is to use division here." but does not explain why division is appropriate or how it will be applied conceptually.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Taylor : Uncorrelated : NA

    Reasoning: Speaker shows their work and explains the strategy's purpose, "See, I made 5 groups of 3. This way, we can multiply 5 by 3 to get the total number of items needed for all 5 days."
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Riley : Explain Own Idea : NA

    Reasoning: Peer A: "The answer is 12." Speaker B: "I disagree with you cause i got 2, 3,... 10." This directly critiques Peer A's conceptual premise (using '10') and provides a correction, making it EOI(HIGH) even though it explains the speaker's own reasoning for the critique.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Speaker B : Engage Others Idea (HIGH) : Peer A
    
    Context:
    - Block Index: {dialogue_info.get("current_block_index", "N/A")}
    - Group: {dialogue_info.get("current_group", "N/A")}
    - Speaker: {dialogue_info.get("current_speaker", "N/A")}
    - Context Description: {dialogue_info.get("context_description", "N/A")}

    Previous Utterances:
    - t-2 ({dialogue_info.get("previous_speaker_2", "N/A")}): {dialogue_info.get("previous_utterance_2", "None")}
    - t-1 ({dialogue_info.get("previous_speaker_1", "N/A")}): {dialogue_info.get("previous_utterance_1", "None")}

    Current Utterance to Classify (ID: {utterance_id_for_output}):
    - Speaker: {dialogue_info.get("current_speaker", "N/A")}
    - Text: "{dialogue_info.get("current_utterance_text", "")}"

    Classification result:
    """

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ]
    return messages

# --- 3. classification ---
def generate_text_classification(provider, api_key, dialogue_info, utterance_id_for_output, model_name=None):
    """Call external LLM provider (OpenAI / Google / Anthropic) to classify a single utterance."""
    messages = make_prompt_messages(dialogue_info=dialogue_info, utterance_id_for_output=utterance_id_for_output)

    # --- 1. Dispatch to provider specific APIs ------------------------------
    full_model_response = ""

    if provider.lower() == "openai":

        client = OpenAI(api_key=api_key)
        if model_name is None:
            model_name = "gpt-4o-mini"  

        # Retry logic for OpenAI API calls
        max_retries = 3
        for attempt in range(max_retries + 1):
            try:
                response = client.chat.completions.create( 
                    model=model_name,
                    messages=messages, 
                    temperature=0,
                )
                
                raw_response_content = getattr(response, "content", None)
                if raw_response_content is None:
                    # Fallback to string representation if content is None
                    full_model_response = str(response).strip()
                else:
                    full_model_response = raw_response_content.strip()

                break  # Success, exit retry loop

            except Exception as e:
                error_str = str(e).lower()
                retryable_errors = ['503', '502', '429', 'timeout', 'service unavailable', 'bad gateway', 'too many requests', 'rate limit', 'attributeerror', 'none']
                is_retryable = any(err in error_str for err in retryable_errors)

                if attempt < max_retries and is_retryable:
                    wait_time = (2 ** attempt) + 1  # Exponential backoff: 2, 5, 9 seconds
                    print(f"  [RETRY] OpenAI API error (attempt {attempt + 1}/{max_retries + 1}): {e}")
                    print(f"  [RETRY] Waiting {wait_time} seconds before retrying...")
                    time.sleep(wait_time)
                else:
                    print(f"  [ERROR] OpenAI API failed after {max_retries + 1} attempts: {e}")
                    raise e

    elif provider.lower() == "anthropic":
        client = anthropic.Anthropic(api_key=api_key)
        if model_name is None:
            model_name = "claude-3-sonnet-20240229"

        # Separate system and user messages for caching
        system_content = None
        user_messages = []
        
        for m in messages:
            if m["role"] == "system":
                system_content = m["content"]
            else:
                user_messages.append({"role": m["role"], "content": m["content"]})

        # Retry logic for Anthropic API calls
        max_retries = 3
        for attempt in range(max_retries + 1):
            try:
                # Use prompt caching for system message
                response = client.messages.create(
                    model=model_name,
                    system=[
                        {
                            "type": "text",
                            "text": system_content,
                            "cache_control": {"type": "ephemeral"}
                        }
                    ] if system_content else None,
                    messages=user_messages,
                    max_tokens=1024,
                    temperature=0,
                )

                # `content` can be list[Block]
                # Added 'attributeerror' to retryable errors
                raw_response_content = "".join(block.text for block in response.content) if response.content else None
                if raw_response_content is None:
                    full_model_response = str(response).strip()
                else:
                    full_model_response = raw_response_content.strip()

                break  # Success, exit retry loop
                
            except Exception as e:
                error_str = str(e).lower()
                # Check for retryable errors (503, 502, 429, timeout, etc.)
                retryable_errors = ['503', '502', '429', 'timeout', 'service unavailable', 'bad gateway', 'too many requests', 'rate limit', 'attributeerror', 'none']
                is_retryable = any(err in error_str for err in retryable_errors)
                
                if attempt < max_retries and is_retryable:
                    wait_time = (2 ** attempt) + 1  # Exponential backoff: 2, 5, 9 seconds
                    print(f"  [RETRY] Anthropic API error (attempt {attempt + 1}/{max_retries + 1}): {e}")
                    print(f"  [RETRY] Waiting {wait_time} seconds before retrying...")
                    time.sleep(wait_time)
                else:
                    # Final attempt failed or non-retryable error
                    print(f"  [ERROR] Anthropic API failed after {max_retries + 1} attempts: {e}")
                    raise e

    elif provider.lower() == "google":
        if model_name is None:
            model_name = "gemini-2.0-flash" 

        system_instruction_content = ""
        user_content = ""
        for m in messages:
            if m["role"] == "system":
                system_instruction_content = m["content"]
            elif m["role"] == "user":
                user_content = m["content"]
        
        client = genai.Client(api_key=api_key)

        max_retries = 5
        for attempt in range(max_retries + 1):
            try:
                response = client.models.generate_content(
                    model=model_name,
                    config=types.GenerateContentConfig(
                        temperature=0,
                        thinking_config=types.ThinkingConfig(thinking_budget=128),
                        system_instruction=system_instruction_content 
                    ),
                    contents=user_content, # user_content를 contents 인자로 전달
                )
                
                if response is None:
                    raise ValueError("Google API returned None response object.")

                raw_response_content = getattr(response, "text", str(response))
                
                if raw_response_content is None:
                    full_model_response = ""
                    print(f"  [WARN] Google API response.text or fallback was None for ID {utterance_id_for_output}. Treating as empty.")
                else:
                    full_model_response = str(raw_response_content).strip()

                break
                
            except Exception as e:
                error_str = str(e).lower()
                retryable_errors = ['500','503', '502', '429', 'timeout', 'service unavailable', 'bad gateway', 'too many requests', 'rate limit', 'attributeerror', 'none']
                is_retryable = any(err in error_str for err in retryable_errors)
                
                if attempt < max_retries and is_retryable:
                    wait_time = (2 ** attempt) + 1
                    print(f"  [RETRY] Google API error (attempt {attempt + 1}/{max_retries + 1}): {e}")
                    print(f"  [RETRY] Waiting {wait_time} seconds before retrying...")
                    time.sleep(wait_time)
                else:
                    print(f"  [ERROR] Google API failed after {max_retries + 1} attempts: {e}")
                    raise e

    else:
        raise ValueError(f"Unsupported provider '{provider}'. Choose from 'openai', 'google', or 'anthropic'.")

    # --- 2. Parse the response to extract classification line ---------------
    print(f"LLM Full Raw Output for ID {utterance_id_for_output}:\n---\n{full_model_response}\n---")

    current_id_str = str(utterance_id_for_output).strip()
    default_error_line = f"{current_id_str} : {dialogue_info.get('current_speaker','ERR_SPEAKER')} : PARSE_ERROR : NA"
    final_output_line = default_error_line
    reasoning_part = ""

    def strip_bf_number(text):
        match = re.search(r'(\d.*)', text)
        return match.group(1) if match else text

    if "FINAL_CLASSIFICATION:" in full_model_response:
        reasoning_part, classification_content_after_marker = full_model_response.split("FINAL_CLASSIFICATION:", 1)
        potential_lines = strip_bf_number(classification_content_after_marker).strip().splitlines()

        found_line = ""
        for candidate in potential_lines:
            cand = candidate.strip()
            if cand.startswith(current_id_str) and cand.count(':') == 3:
                found_line = cand
                break
            elif not found_line and cand:  # backup if strict match not found
                found_line = cand

        if found_line:
            cleaned = found_line
            match = re.search(rf"({re.escape(current_id_str)}\s*:.*)", cleaned)
            if match:
                cleaned = match.group(1)
            cleaned = re.sub(r"^[*\s#\-:!]+", "", cleaned)

            if cleaned.startswith(f"{current_id_str} :") and cleaned.count(':') == 3:
                final_output_line = cleaned
            else:
                print(f"  [WARN] Parsed line format mismatch for ID {current_id_str}. Raw: '{found_line}'")
        else:
            print(f"  [WARN] 'FINAL_CLASSIFICATION:' found but no valid line for ID {current_id_str}.")

    else:
        # Fallback: search last lines for id pattern
        for line in reversed(full_model_response.splitlines()):
            line = line.strip()
            if not line:
                continue
            candidate = re.sub(r"^[*\s#\-:!]+", "", line)
            if candidate.startswith(f"{current_id_str} :") and candidate.count(':') == 3:
                final_output_line = candidate
                break
        
        lines = full_model_response.splitlines()
        if lines:
            reasoning_part = "\n".join(lines[:-1]).strip()
        else:
            reasoning_part = ""

    return final_output_line, reasoning_part.strip()

# --- 4. main function ---
def process_dialogue_dataframe(df, provider, api_key, model_name=None, sleep_sec: float = 0.0):
    results = []
    
    if 'timestamp' in df.columns and not pd.api.types.is_string_dtype(df['timestamp']):
        try:
            df['timestamp_str'] = df['timestamp'].astype(str) 
        except Exception as e:
            print(f"Warning: Could not convert 'timestamp' column to string: {e}. Assuming it's already string-like or will be handled by parser.")
            df['timestamp_str'] = df['timestamp'] 
    elif 'timestamp' in df.columns:
         df['timestamp_str'] = df['timestamp']
    else:
        print("Warning: 'timestamp' column not found. Timestamps will be 'N/A'.")
        df['timestamp_str'] = "N/A"

    def timestamp_to_seconds(ts_str):
        if pd.isna(ts_str) or ts_str == "N/A" or not isinstance(ts_str, str): return -1
        parts = str(ts_str).split(':')
        try:
            if len(parts) == 2: return int(parts[0])*60 + int(parts[1])
            if len(parts) == 3: return int(parts[0])*3600 + int(parts[1])*60 + int(parts[2])
        except ValueError:
            return -2 
        return 0 
    
    df['sortable_time'] = df['timestamp_str'].apply(timestamp_to_seconds)    
    df['utterance_id'] = df.index.astype(str)

    current_block_index = 0
    last_timestamp_str, last_group, last_date_str = None, None, None
    previous_utterances_in_block = []
    
    for index, row in tqdm(df.iterrows(), total=len(df), desc="Processing Dialogues"):
        current_utterance_id = str(row['utterance_id']) 
        current_group = str(row.get('group', "UnknownGroup"))
        current_date_str = str(row.get('date', "UnknownDate")) 
        current_timestamp_str = str(row.get('timestamp_str', "N/A"))
        
        current_speaker = str(row.get('name', "UnknownSpeaker"))
        current_utterance_text = str(row.get('utterance', ""))

        if (last_group != current_group or 
            last_date_str != current_date_str or 
            last_timestamp_str != current_timestamp_str):
            current_block_index += 1
            previous_utterances_in_block = []

        dialogue_info = {
            "current_block_index": current_block_index,
            "current_group": current_group,
            "current_date": current_date_str, 
            "current_speaker": current_speaker,
            "current_utterance_text": current_utterance_text,
            "current_utterance_indent": 0
        }
        if len(previous_utterances_in_block) >= 1:
            dialogue_info["previous_utterance_1"] = previous_utterances_in_block[-1]['utterance']
            dialogue_info["previous_speaker_1"] = previous_utterances_in_block[-1]['speaker']
        if len(previous_utterances_in_block) >= 2:
            dialogue_info["previous_utterance_2"] = previous_utterances_in_block[-2]['utterance']
            dialogue_info["previous_speaker_2"] = previous_utterances_in_block[-2]['speaker']
        if len(previous_utterances_in_block) >= 3: 
            dialogue_info["previous_utterance_3"] = previous_utterances_in_block[-3]['utterance']
            dialogue_info["previous_speaker_3"] = previous_utterances_in_block[-3]['speaker']

        classification_result_line = f"{current_utterance_id} : {current_speaker} : CLASSIFICATION_FAILED : NA"
        reasoning_part = ""
        if not current_utterance_text.strip():
            classification_result_line = f"{current_utterance_id} : {current_speaker} : Uncorrelated : NA"
            reasoning_part = ""
        else:
            try:
                classification_result_line, reasoning_part = generate_text_classification(
                    provider, api_key, dialogue_info, current_utterance_id, model_name=model_name
                )
                if sleep_sec > 0:
                    time.sleep(sleep_sec)
            except Exception as e:
                print(f"  [CRITICAL ERROR] Classification failed for utterance ID {current_utterance_id}: '{current_utterance_text[:50]}...' - {e}")
                import traceback; traceback.print_exc()
                reasoning_part = ""
        results.append({
            'utterance_id': current_utterance_id,
            'timestamp': current_timestamp_str,
            'name': current_speaker,
            'utterance': current_utterance_text,
            'group': current_group,
            'date': current_date_str,
            'cls_original': row.get('cls', 'N/A'),
            'block_index': current_block_index,
            'llm_classification_output_line': classification_result_line,
            'llm_reasoning': reasoning_part,
        })

        previous_utterances_in_block.append({'speaker': current_speaker, 'utterance': current_utterance_text})
        if len(previous_utterances_in_block) > 3:
            previous_utterances_in_block.pop(0)
        
        last_timestamp_str, last_group, last_date_str = current_timestamp_str, current_group, current_date_str
            
        if (index + 1) % 20 == 0:
            pass  
    
    return pd.DataFrame(results)

# --- 5. LLM output parsing ---
def parse_llm_output_line(line_str, expected_id_str):
    parts = line_str.split(' : ')
    if len(parts) == 4:
        parsed_id = parts[0].strip()
        parsed_name = parts[1].strip()
        llm_label_full = parts[2].strip() 
        llm_reference = parts[3].strip()

        llm_label_main = llm_label_full
        llm_engagement = "NA" 

        if "Engage Others Idea" in llm_label_full:
            if "(LOW)" in llm_label_full:
                llm_label_main = "Engage Others Idea"
                llm_engagement = "LOW"
            elif "(MEDIUM)" in llm_label_full:
                llm_label_main = "Engage Others Idea"
                llm_engagement = "MEDIUM"
            elif "(HIGH)" in llm_label_full:
                llm_label_main = "Engage Others Idea"
                llm_engagement = "HIGH"

        result_dict = {
            "parsed_utterance_id": parsed_id,
            "parsed_name": parsed_name,
            "parsed_indent": "0",  # Default value since it's no longer in the input
            "llm_label": llm_label_main,      
            "llm_engagement": llm_engagement, 
            "llm_reference": llm_reference
        }

        if parsed_id == expected_id_str:
            return result_dict
        else:
            print(f"  [PARSE_ID_MISMATCH] Expected ID {expected_id_str} but got {parsed_id} in line: '{line_str}'")
            result_dict["llm_label"] = f"ID_MISMATCH_{llm_label_main}"
            return result_dict

    print(f"  [PARSE_WARN] LLM output line does not have 4 parts. Line: '{line_str}' for ID {expected_id_str} (Found {len(parts)} parts)")
    return {
        "parsed_utterance_id": str(expected_id_str), 
        "parsed_name": "ERR_FORMAT",
        "parsed_indent": "0",  # Default value
        "llm_label": "LINE_FORMAT_ERROR", 
        "llm_engagement": "NA",
        "llm_reference": "NA"
    }

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description="Process dialogue transcripts for utterance classification via external LLM APIs.")
    parser.add_argument("--provider", choices=["openai", "google", "anthropic"], required=True, help="LLM provider to use.")
    parser.add_argument("--apikey_json", type=str, default="apikey.json", help="Path to apikey.json containing provider keys.")
    parser.add_argument("--model_name", type=str, default=None, help="Override the default model name for the chosen provider.")
    parser.add_argument("--transcript_dir", type=str, default="./transcripts/", help="Directory with .docx transcripts.")
    parser.add_argument("--output_csv", type=str, default="classification_results.csv", help="Output CSV filename.")
    parser.add_argument("--sleep_sec", type=float, default=0.0, help="Seconds to sleep between consecutive API calls (rate-limit throttle).")

    args = parser.parse_args()

    start_time = time.time()
    print(f"Provider: {args.provider}")
    print(f"Transcript directory: {args.transcript_dir}")
    print(f"Output CSV: {args.output_csv}")

    # --- Load API Key -------------------------------------------------------
    try:
        with open(args.apikey_json, "r", encoding="utf-8") as f:
            key_map = json.load(f)
        api_key = key_map.get(f"{args.provider}_api_key")
        if not api_key:
            raise ValueError(f"API key for provider '{args.provider}' not found in {args.apikey_json}")

        opath = args.transcript_dir
        os.makedirs(opath, exist_ok=True)
        if not os.listdir(opath) and opath == "./test_transcripts/":
            print(f"Creating dummy docx files in {opath} for testing...")
            dummy_doc_content1="""0:01\nSpeakerA: Hello.\n0:02\nSpeakerB: Hi?\n0:03\nSpeakerA: Add 2 and 2.\n0:05\nSpeakerC: 4.\n0:06\nSpeakerA: Yes."""
            doc1=docx.Document(); [doc1.add_paragraph(l.strip()) for l in dummy_doc_content1.strip().split('\n')]; doc1.save(os.path.join(opath,"AL_G1_Test_1.1_f.docx"))
            print("Dummy file created.")

        fpath = sorted([f for f in os.listdir(opath) if f.endswith(".docx") and not f.startswith("~")])
        if not fpath: print(f"No .docx files found in {opath}. Exiting."); exit()
        print(f"Found {len(fpath)} docx files.")

        group_list = ["total" if ("ipad"in f.lower())or("finalized"in f.lower()) else f.split("_")[1] if len(f.split("_"))>1 else "grp_err" for f in fpath]
        date_list = [re.findall(r'\d{1,2}\.\d{1,2}',f)[0] if re.findall(r'\d{1,2}\.\d{1,2}',f) else "date_err" for f in fpath]
        class_list = [f[0:2] if len(f)>=2 else "cls_err" for f in fpath]
        
        full_path = pd.DataFrame({"group":group_list,"date":date_list,"cls":class_list,"path":[os.path.join(opath,f) for f in fpath]})
        full_path['group'] = full_path['group'].str.replace(")", "", regex=False).str.replace("(", "_", regex=False)
        
        full_path_filtered = full_path[(full_path['cls'] == "AL") & (full_path['group'] != "total")].reset_index(drop=True)
        N = len(full_path_filtered)
        if N == 0: print("No files after filtering. Exiting."); exit()
        print(f"Processing {N} files after filtering.")

        data_list = []
        for _, row_fp in tqdm(full_path_filtered.iterrows(), total=N, desc="Extracting Docx"):
            try: data_list.append(extract_dialogue_from_docx(row_fp['path']).assign(group=row_fp['group'],date=row_fp['date'],cls=row_fp['cls']))
            except Exception as e: print(f"Error extracting {row_fp['path']}: {e}")
        
        if not data_list: print("No data extracted. Exiting."); exit()
        df_input = pd.concat(data_list, ignore_index=True)
        if df_input.empty: print("Concat DataFrame empty. Exiting."); exit()
        print(f"\nTotal utterances to process: {len(df_input)}")

        df_results = process_dialogue_dataframe(df_input.copy(), args.provider, api_key, model_name=args.model_name, sleep_sec=args.sleep_sec)

        print("\n--- Processing Complete ---")

        parsed_data = [ {**row_res.to_dict(), **parse_llm_output_line(str(row_res['llm_classification_output_line']), str(row_res['utterance_id']))} for _, row_res in df_results.iterrows()]
        df_final_results = pd.DataFrame(parsed_data)
        df_final_results.drop(columns=['utterance', 'reason'], errors='ignore', inplace=True)
        df_final_results.dropna(how='all', axis=1, inplace=True)
        df_final_results.to_csv(args.output_csv, index=False, encoding='utf-8-sig') 
        print(f"\nResults saved to {args.output_csv}")

        total_time = time.time() - start_time
        print(f"Total execution time: {total_time:.2f}s ({total_time/60:.2f}min).")
        if len(df_input)>0: print(f"Avg time/utterance: {total_time/len(df_input):.2f}s.")

    except Exception as e:
        print(f"A critical error occurred in main execution: {e}")
        import traceback; traceback.print_exc()