import os
import pandas as pd
import torch
from transformers import AutoTokenizer, AutoModelForCausalLM, AutoConfig, GenerationConfig, BitsAndBytesConfig, AutoProcessor
import docx
import re
from tqdm import tqdm
import time
import argparse

"""
python llm.py --model_id "[model name]" --output_csv "[output file name]" --transcript_dir "[data path]"
"""

# --- 1. Docx parser ---
def extract_dialogue_from_docx(docx_path):
    doc = docx.Document(docx_path)
    full_text = [paragraph.text for paragraph in doc.paragraphs]
    text = '\n'.join(full_text)

    chunks = re.split(
        r'(\n+\d{1,2}:\d{2}(?::\d{2})?\s*\n|\n+\d{1,2}:\d{2}(?::\d{2})?\s*$|^\d{1,2}:\d{2}(?::\d{2})?\s*\n)',
        text, flags=re.MULTILINE
    )

    timestamps, speakers, utterances = [], [], []
    current_timestamp = None

    for chunk in chunks:
        clean_chunk = chunk.strip('\n').strip()
        if re.fullmatch(r"\d{1,2}:\d{2}(?::\d{2})?", clean_chunk):
            current_timestamp = clean_chunk
        else:
            lines = [l.strip() for l in chunk.split('\n') if l.strip()]
            i = 0
            while i < len(lines):
                line = lines[i]
                match = re.match(r'^([A-Za-z가-힣]+):\s*(.+)', line)
                if match:
                    speaker, utterance = match.group(1), match.group(2)
                    utterance = re.sub(r'\s+', ' ', utterance)
                    timestamps.append(current_timestamp)
                    speakers.append(speaker)
                    utterances.append(utterance)
                elif re.match(r'^\d{1,2}:\d{2}(?::\d{2})?\s+\[.+\]$', line):
                    ts, desc = re.match(r'^(\d{1,2}:\d{2}(?::\d{2})?)\s+(\[.+\])$', line).groups()
                    current_timestamp = ts
                    timestamps.append(current_timestamp)
                    speakers.append("NARRATION")
                    utterances.append(desc)
                elif re.match(r'^\[.+\]$', line):
                    timestamps.append(current_timestamp)
                    speakers.append("NARRATION")
                    utterances.append(line)
                i += 1
    df = pd.DataFrame({'timestamp': timestamps, 'name': speakers, 'utterance': utterances})
    df = df[df.name != "NARRATION"].reset_index(drop=True)
    return df



# --- 2. Model load ---
def load_model_and_tokenizer(model_id):
    model_id_lower = model_id.lower()    
    print(f"Loading tokenizer for {model_id}...")
    tokenizer = AutoTokenizer.from_pretrained(model_id, trust_remote_code=True)        
    if tokenizer.model_max_length is None or tokenizer.model_max_length > 4096 * 128:
        tokenizer.model_max_length = 4096 * 128
        print(f"Adjusted other model_max_length to {tokenizer.model_max_length}")
    
    if tokenizer.pad_token is None:
        if tokenizer.eos_token is not None:
            tokenizer.pad_token = tokenizer.eos_token
            print(f"Tokenizer pad_token was None, set to eos_token: '{tokenizer.pad_token}' (ID: {tokenizer.eos_token_id})")
        else:
            tokenizer.add_special_tokens({'pad_token': '[PAD]'})
            print(f"Tokenizer pad_token and eos_token were None. Added pad_token: {tokenizer.pad_token} (ID: {tokenizer.pad_token_id})")

    print(f"Loading model {model_id}...")
    model_kwargs = {
        "device_map": "auto",
        "trust_remote_code": True
    }

    if "llama-3" in model_id_lower:
        model_kwargs["torch_dtype"] = torch.bfloat16 if torch.cuda.is_bf16_supported() else torch.float16
    elif "phi" in model_id_lower: 
        model_kwargs["torch_dtype"] = "auto"
    elif "qwen" in model_id_lower:
        model_kwargs["torch_dtype"] = torch.bfloat16 if torch.cuda.is_bf16_supported() else torch.float16
    elif "gemma" in model_id_lower:
        print(f"Applying optimized settings for Gemma model: {model_id}")
        target_dtype = torch.bfloat16 if torch.cuda.is_bf16_supported() else torch.float16
        quant_config = BitsAndBytesConfig(
                load_in_4bit=True,
                bnb_4bit_quant_type="nf4",
                bnb_4bit_compute_dtype=target_dtype,
                bnb_4bit_use_double_quant=True,
        )
        
        model_kwargs = {
            "device_map": "auto",                
            "trust_remote_code": True,
            "low_cpu_mem_usage": True,           
            "quantization_config": quant_config, 
            "torch_dtype": target_dtype,         
        }

    else:
        model_kwargs["torch_dtype"] = torch.float16
    
    if "phi" in model_id_lower or "qwen" in model_id_lower:
        try:                
            print(f"Attempting to set 'flash_attention_2' for {model_id}")
            model_kwargs["attn_implementation"] = "flash_attention_2"
            print(f"{model_id}: Set to use 'flash_attention_2'.")
        except ImportError:
            print(f"{model_id}: 'flash_attention_2' setup failed. Using default.")
            model_kwargs["attn_implementation"] = "sdpa" 
    
    ModelClass = AutoModelForCausalLM
    try:
        print(f"Loading model '{model_id}' with kwargs: {model_kwargs}")
        model = ModelClass.from_pretrained(model_id, **model_kwargs)
        print(f"{model_id} loaded successfully.")
    except Exception as e1:
        print(f"Error loading model {model_id} with primary kwargs: {e1}")
        if "attn_implementation" in model_kwargs and model_kwargs["attn_implementation"] != "sdpa":
            original_attn = model_kwargs.pop("attn_implementation")
            print(f"Retrying without '{original_attn}', falling back to 'sdpa'...")
            model_kwargs["attn_implementation"] = "sdpa" 
            try:
                model = ModelClass.from_pretrained(model_id, **model_kwargs)
                print(f"{model_id} loaded successfully on retry with sdpa.")
            except Exception as e2:
                print(f"Retry with 'sdpa' failed: {e2}")
                print("Trying with minimal kwargs...")
                minimal_kwargs = {
                    "trust_remote_code": True,
                    "device_map": "auto",
                }
                
                if "torch_dtype" in model_kwargs:
                    minimal_kwargs["torch_dtype"] = model_kwargs["torch_dtype"]
                if "quantization_config" in model_kwargs:
                    minimal_kwargs["quantization_config"] = model_kwargs["quantization_config"]
                if "low_cpu_mem_usage" in model_kwargs:
                    minimal_kwargs["low_cpu_mem_usage"] = model_kwargs["low_cpu_mem_usage"]

                try:
                    print(f"Retrying with minimal_kwargs: {minimal_kwargs}")
                    model = ModelClass.from_pretrained(model_id, **minimal_kwargs)
                    print(f"{model_id} loaded successfully with minimal kwargs.")
                except Exception as e_minimal:
                    print(f"CRITICAL: Model loading failed with all fallbacks for {model_id}: {e_minimal}")
                    raise
        else:
            print(f"CRITICAL: Model loading failed, no attention implementation fallbacks for {model_id}.")
            raise e1
    
    if hasattr(model.config, 'pad_token_id') and model.config.pad_token_id is None and tokenizer.pad_token_id is not None:
        model.config.pad_token_id = tokenizer.pad_token_id
        print(f"Model pad_token_id was None, set to tokenizer's pad_token_id: {model.config.pad_token_id}")
       
    if not getattr(model.config, "chat_template", None) and not getattr(tokenizer, "chat_template", None):        
        print(f"[WARN] Model {model_id}: No chat_template found. Manual prompting may be needed.")      

    model.eval()
    if hasattr(model.config, 'use_cache'):
        model.config.use_cache = True    
    print(f"Model '{model_id}' and its tokenizer are ready.")
    return model, tokenizer


# --- 3. Prompt ---
def make_prompt_messages(dialogue_info, utterance_id_for_output):
    system_prompt = f"""
    You are an expert in analyzing utterance contributions in problem-solving discussions. Your task is to perform two steps:

    IMPORTANT NOTE ABOUT UTTERANCE CONTENT:
    - Content within parentheses () or square brackets [] represents ACTIONS/BEHAVIORS, not speech.
    - Example: "[points to the board]", "(writes on paper)", "[moves blocks]" are actions, not verbal utterances.
    - If an utterance consists entirely of action descriptions with no speech, classify it as Uncorrelated (UC).

    STEP 1: REASONING
    Analyze the utterance based on the rules below. BRIEFLY explain your reasoning (1-2 concise sentences) on its contribution to idea development.
    Example: Reasoning: Speaker explains their method of adding 10s, fitting EXP criteria as it details their reasoning.

    STEP 2: CLASSIFICATION OUTPUT
    Provide the classification as a SINGLE line: "FINAL_CLASSIFICATION: [UTTERANCE_ID] : [NAME] : [LABEL] : [REFERENCE]"
    - EXP & UC: [REFERENCE] is 'NA'.
    - EOI: [REFERENCE] is the NAME of the participant whose idea is engaged (trace back if needed). Use 'Previous Speaker' if name is unknown but clearly refers to t-1 idea.

    --- DECISION FLOW ---
    Follow this order strictly:
    1.  Check for EOI (Engage Others Idea) first. Does the utterance DIRECTLY respond to, critique, correct, or question a specific, conceptual idea from a peer? If YES, classify as EOI and determine the level (LOW, MEDIUM, HIGH).
    2.  If not EOI, check for EXP (Explain Own Idea). Is the speaker verbalizing their own conceptual reasoning, strategy, or the 'why' behind their actions? If YES, classify as EXP.
    3.  If not EOI or EXP, classify as UC (Uncorrelated). This is the default for anything that lacks clear conceptual contribution.

    --- CRITICAL WARNING: AVOID COMMON MISTAKES ---
    - UC vs. EOI(LOW): A simple "Yes", "No", "I agree", or "Wow!" is NOT EOI. It is Uncorrelated (UC) unless it clearly rephrases the peer's conceptual idea, proving understanding. Social agreement alone is not conceptual engagement.
    - UC vs. EXP: Simply describing an action ("I'm counting") or stating an answer ("It's 35") is Uncorrelated (UC). To be EXP, the speaker must explain the purpose of the action ("I'm counting to check my answer") or the reasoning behind the answer ("It's 35 because 70 divided by 2 is 35").

    --- CATEGORY DEFINITIONS & CRITERIA ---

    1.  Explain Own Idea (EXP):
        -   Definition: Speaker verbalizes their own conceptual reasoning, mathematical strategy, or understanding with clear explanation, justification, or elaboration that reveals their thought process. This includes explanations of the speaker's own ideas or actions, even if prompted by a peer's question.
        -   Key Includes:
            *   Elaborating a novel strategy or a significant modification to an existing one.
            *   Explaining a conceptual realization with its basis.
            *   Justifying a mathematical method or answer.
            *   Articulating a step-by-step conceptual plan or calculation process with its underlying logic.

    2.  Engage Others Idea (EOI): Speaker directly engages with the content of another's specific articulated conceptual mathematical idea. The level depends on the depth and nature of this engagement.

        -   EOI(LOW): Minimum Conceptual Engagement - Confirmation & Basic Clarification
            *   Definition: Speaker briefly acknowledges, confirms, agrees/disagrees with, or asks a very simple, direct question about a clearly identifiable, specific conceptual mathematical point previously articulated by a peer. The utterance MUST demonstrate BOTH (1) clear understanding of that specific conceptual point AND (2) make a direct, unambiguous verbal reference to it.
            *   Key Indicators: Confirmatory Restatement, Targeted Agreement/Disagree-ment, Simple clarification on a specific detail within the peer's idea.

        -   EOI(MEDIUM): Probing & Minor Elaboration (of a peer's idea)
            *   Definition: Speaker seeks to deepen their understanding of a peer's specific conceptual idea by asking questions that probe the underlying reasoning, assumptions, or implications of that peer's idea. They might also restate the peer's idea with minor conceptual additions or offer small, scaffolded extensions.
            *   Key Indicators: Probing "Why" or "How" Questions about a peer's reasoning, Requesting Elaboration on a peer's idea, Identifying Ambiguity for clarification within a peer's idea.

        -   EOI(HIGH): Critique, Significant Extension, & Synthesis (of a peer's idea)
            *   Definition: Speaker critically evaluates a peer's specific conceptual idea by offering a reasoned critique, proposing a significant and justified alternative/refinement to that peer's idea, or by synthesizing it with other ideas to create a more robust solution. This involves substantial conceptual work.
            *   Clarification on Critique: A critique must be reasoned. A simple, unelaborated disagreement or correction (e.g., 'No, the answer is 20') does not qualify as EOI(HIGH). It is a targeted disagreement about a conceptual point but lacks justification, making it EOI(LOW). To be EOI(HIGH), the utterance must explain why the peer's idea is incorrect (e.g., 'No, it's 20, because you forgot to add the 2.'), or clearly justify their disagreement by demonstrating their own correct reasoning or solution process (verbally or through action) in a manner that implicitly highlights the peer's error in relation to their original idea.
            *   Key Indicators: Reasoned Conceptual Critique of a peer's idea, Correction with Justification related to a peer's error, Proposing Justified Alternatives to a peer's idea, Conceptual Synthesis of a peer's idea with others.
            *   Proposing a Conceptual Paradigm Shift: This applies when a speaker clearly rejects a peer's more concrete, lower-level strategy and proposes a shift to a much more abstract or higher-level conceptual category. In this case, even if the proposed alternative does not include a specific methodology, the 'leap' itself is considered a significant critique and proposal that fundamentally changes the direction of problem-solving.

    3.  Uncorrelated (UC):
        -   Definition: No meaningful or discernible conceptual mathematical contribution; utterance is vague, off-task, purely procedural without explanation, or a brief unelaborated realization/intention.
        -   Key Includes:
            *   Narration of one's own actions without explaining the conceptual purpose.
            *   Vague statements ("This is hard," "I'm confused" without specifying the conceptual point).
            *   Simple affirmations/negations ("Yes", "I agree", "Okay") without specific reference to a peer's conceptual idea and demonstrating understanding of it.
            *   Generic procedural questions/statements ("What's next?").
            *   Unelaborated realizations ("I get it!", "I know!", "Oh!").
            *   Stating answers or numbers without explaining the process or reasoning.
            *   Reading aloud from a worksheet or simply repeating a peer's words without adding conceptual value.
            *   Brief self-corrections ('Oops') without explaining the conceptual reason for the correction.

    --- CORE CLASSIFICATION PRINCIPLES ---
    1. Conceptual Contribution Focus: Your primary focus is on conceptual idea development in mathematics.
    2. EOI - Engaging a Peer's Conceptual Idea: EOI requires engagement with a specific peer's conceptual mathematical idea, not just their words or a surface-level answer.
    3. EXP - Explaining Own Conceptual Reasoning: EXP requires explaining one's own reasoning, strategy, or understanding with clear justification or elaboration.
    4. Priority for Primary Conceptual Function: Choose the category that best reflects the utterance's dominant conceptual purpose. In complex cases where functions are mixed, use the Tie-Breaker rules below to help determine this primary function.
    5. UC as Default, with Caution for EOI(LOW): If an utterance does not clearly meet the criteria for EXP or EOI, classify it as UC. However, do not default to UC too aggressively if it meets the minimum criteria for EOI(LOW).
    6. Context is Key: Your classification must be based on the flow of the dialogue. Look beyond the immediate previous utterance; trace back to the origin of an idea for accurate EOI referencing or to understand the full context for an EXP.
    7. The "Target of Reasoning" Principle (EOI Precedence Rule): If an utterance's primary conceptual function is to respond directly to, critique, correct, or build upon a specific peer's articulated idea, it MUST be classified as EOI.
        - This is a strict priority rule. It holds true even if the speaker explains their own reasoning or calculation process as the means for that critique or extension. The target of the idea (a peer) takes precedence over the form of the reasoning (one's own calculation).
        - Critical Example (Correction as EOI(HIGH)):
            Peer A: "I got 18."
            Speaker B: "No, 9 plus 9 is 18, but you forgot to add the 2, so it's 20."
            Correct Classification: EOI(HIGH) : Peer A. This is NOT EXP. Although Speaker B explains their own calculation, the entire purpose of the utterance is to correct Peer A's idea.
    8. Principle of Conceptual Impact Over Explicit Explanation:
        - Do not automatically classify an utterance as UC just because it lacks an explicit "because" clause. Evaluate its function within the dialogue.
        - The Test: Does the utterance introduce a new, relevant mathematical object, operation, or logical step into the conversation? If yes, it likely has conceptual value.
            - Lean towards UC if it's just stating a number without context. (e.g., Peer A: "This is hard." Speaker B: "70.")
            - Lean towards EXP or EOI if it introduces a key operation or a logical next step based on the current context. (e.g., After the group agrees the total is 70 -> Speaker B: "Okay, so 70 divided by 2.") This utterance, despite its brevity, introduces the crucial next conceptual step (division) and is therefore not UC.
    9. Principle of Utterance Continuity: If an utterance is a direct continuation of the speaker's own previous turn (e.g., finishing a sentence or elaborating on a point immediately after pausing), evaluate the current utterance in the context of the combined statement. The classification should reflect the conceptual contribution of the complete idea, not just the final fragment.

    --- TIE-BREAKER RULES FOR BORDERLINE CASES ---
    The following are priority principles to increase consistency when classifying ambiguous, borderline cases.

    1. (EOI vs. EXP) The "Function over Form" Rule (Highest Priority):
    - This rule takes precedence over all others when an utterance has mixed characteristics. It applies when an utterance contains the speaker's own reasoning or calculation (form of EXP) but is used to directly critique, correct, or offer a counter-proposal to a peer's specific idea (function of EOI).
    - The Test: Ask, "What is the primary purpose of this utterance in the dialogue?"
        - If the purpose is to correct a peer's error, challenge a peer's method, or propose an alternative in direct response to a peer, it MUST be classified as EOI(HIGH). The corrective/critical function outweighs the explanatory form.
    - Critical Failure Case to Avoid:
        - Peer A: "I got the wrong answer, 18."
        - Speaker B: "No, because 9 plus 9 is 18, but you forgot to add the 2, so the answer is 20."
        - Correct Classification: Engage Others Idea (HIGH) : Peer A.
        - Incorrect Classification: Explain Own Idea : NA. Do not classify this as EXP simply because Speaker B explains their own calculation. The entire utterance exists only because of Peer A's mistake and its primary function is to correct it.

    2. (UC vs. EOI(LOW)) The "Rephrasing vs. Repeating" Rule:
    - An utterance that simply repeats a key word (e.g., "multiplication") or number from a peer should lean towards UC. This is because it provides insufficient evidence of conceptual understanding.
    - Conversely, classify as EOI(LOW) if the speaker rephrases the peer's idea in their own words (even slightly) or includes a specific result based on that idea (e.g., "Oh, so you mean 3 times 4 is 12"). This proves the speaker actively processed the idea.
    - The Test: Did the speaker parrot the peer's idea, or did they digest and explain it in their own head? Only the latter qualifies as EOI(LOW).

    3. (UC vs. EXP) The "Describing Actions vs. Explaining Purpose" Rule:
    - Simply describing an action ("I'm grouping these by 10s") is UC.
    - Explaining the conceptual purpose or reason (the 'Why') behind the action makes it EXP (e.g., "I'm grouping these by 10s to make them easier to count later").
    - Clarification on "Purpose": Explaining how a mathematical strategy is being constructed or how one is modeling the problem's components is also considered "explaining the purpose." The "why" is inherent in the strategic choice of representation.
        - Example of Conceptual Modeling (EXP): "So I'm putting 70 blocks here, and I'm making two groups, to show the total for the two days."
        - This is EXP because it explains how the actions (making groups, using 70 blocks) map onto the concepts of the problem (two days, total).

    4. (EXP vs. EOI(HIGH)) The "Foundation vs. Spark" Rule (for judging idea ownership):
    - To decide, ask this question: "Without the peer's idea, could the speaker's idea still stand on its own?"
    - EOI(HIGH) (Foundation Principle): This applies when the peer's idea serves as the 'foundation' for the speaker's idea. The speaker is modifying, critiquing, or extending the peer's idea, and without it, the speaker's utterance would lose its meaning.
    - EXP (Spark Principle): This applies when the peer's idea only served as a 'spark' for the speaker to come up with a conceptually independent idea. The speaker's idea would still stand as a complete strategy on its own.

    5. (EOI(LOW) vs. EOI(MEDIUM)) The "New vs. Implied Information" Test:
    - Lean towards EOI(LOW): if the utterance primarily confirms or rephrases information that was already stated or clearly implied by the peer. It demonstrates understanding of the existing idea. (e.g., Peer: "We need to split 70 in half." Speaker: "Oh, so 70 divided by 2.")
    - Lean towards EOI(MEDIUM): if the utterance adds a new, explicit piece of conceptual information (like a label, context, or reason) that was not present in the peer's original statement. It enriches the existing idea. (e.g., Peer: "We need to use the 2." Speaker: "Right, because the 2 represents the days.")
    """

    user_prompt = f"""
    Your task:
    Using the rules from the system prompt, classify the utterance with ID {utterance_id_for_output}. First give a brief reasoning, then provide a single line classification in the required format.

    Your response must include:
    1. A concise reasoning (STEP 1)
    2. One strict "FINAL_CLASSIFICATION:" line (STEP 2)
    
    Output format examples:
    Reasoning: Speaker explains own detailed method for grouping and multiplying.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Julian : Explain Own Idea : NA

    Reasoning: Speaker asks "Why did you decide to multiply there?" about John's specific step in his solution.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Julian : Engage Others Idea (MEDIUM) : John

    Reasoning: Speaker is just counting aloud "1, 2, 3" while moving blocks, no explanation of the conceptual plan.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Delilah : Uncorrelated : NA

    Reasoning: Peer A: "I think the pattern is adding 3 each time." Speaker B: "Yes, I see that too, it's adding 3." This shows understanding of and agreement with Peer A's specific conceptual pattern.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Keila : Engage Others Idea (LOW) : Peer A

    Reasoning: Peer A: "The answer is 12." Speaker B: "I agree." This is a generic agreement without referencing the conceptual method or reasoning of Peer A.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Keila : Uncorrelated : NA   

    Reasoning: Peer A states, "The result is 45." Speaker B says, "No, it's 50." This is a direct refutation of a peer's conceptual point (the result) but provides no reasoning or justification, qualifying it as low-level engagement.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Kenji : Engage Others Idea (LOW) : Peer A

    Reasoning: Speaker's utterance "...so we divide by two" completes their thought from their previous utterance "To find the amount for one day...". The combined idea explains the strategy of dividing by two to find the daily amount, making it EXP.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Sarah : Explain Own Idea : NA

    Reasoning: Speaker states "I know how to solve this now!" but doesn't explain the method or what they realized conceptually.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Alex : Uncorrelated : NA

    Reasoning: Speaker critiques Chris's method: "Your addition is correct, but multiplying 5 by 3 would be a more direct conceptual approach here."
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Ben : Engage Others Idea (HIGH) : Chris

    Reasoning: Speaker says "Okay, let's start problem 2." This is procedural coordination.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Maria : Uncorrelated : NA

    Reasoning: Speaker asks a standalone "What?" after a peer's long explanation, without indicating what part is unclear.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : David : Uncorrelated : NA

    Reasoning: Peer C: "So we need to find the total area." Speaker D: "So, by 'total area', you mean we should multiply length by width for each rectangle?" This clarifies understanding of the peer's specific conceptual goal.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Chloe : Engage Others Idea (LOW) : Peer C

    Reasoning: Speaker explains their realization: "Ah, now I understand! 'Each' means we need to perform the operation for every single item, not just once for the total." This explains what was understood conceptually.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Sam : Explain Own Idea : NA

    Reasoning: Speaker says "My strategy is to use division here." but does not explain why division is appropriate or how it will be applied conceptually.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Taylor : Uncorrelated : NA

    Reasoning: Speaker shows their work and explains the strategy's purpose, "See, I made 5 groups of 3. This way, we can multiply 5 by 3 to get the total number of items needed for all 5 days."
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Riley : Explain Own Idea : NA

    Reasoning: Peer A: "The answer is 12." Speaker B: "I disagree with you cause i got 2, 3,... 10." This directly critiques Peer A's conceptual premise (using '10') and provides a correction, making it EOI(HIGH) even though it explains the speaker's own reasoning for the critique.
    FINAL_CLASSIFICATION: {utterance_id_for_output} : Speaker B : Engage Others Idea (HIGH) : Peer A
    
    Context:
    - Block Index: {dialogue_info.get("current_block_index", "N/A")}
    - Group: {dialogue_info.get("current_group", "N/A")}
    - Speaker: {dialogue_info.get("current_speaker", "N/A")}
    - Context Description: {dialogue_info.get("context_description", "N/A")}

    Previous Utterances:
    - t-2 ({dialogue_info.get("previous_speaker_2", "N/A")}): {dialogue_info.get("previous_utterance_2", "None")}
    - t-1 ({dialogue_info.get("previous_speaker_1", "N/A")}): {dialogue_info.get("previous_utterance_1", "None")}

    Current Utterance to Classify (ID: {utterance_id_for_output}):
    - Speaker: {dialogue_info.get("current_speaker", "N/A")}
    - Text: "{dialogue_info.get("current_utterance_text", "")}"

    Classification result:
    """

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ]
    return messages

# --- 4. classification ---
@torch.no_grad()
def generate_text_classification(model, tokenizer, dialogue_info, utterance_id_for_output, model_id="unknown"):
    messages = make_prompt_messages(dialogue_info=dialogue_info, utterance_id_for_output=utterance_id_for_output)

    try:
        prompt_for_model = tokenizer.apply_chat_template(
            messages, tokenize=False, add_generation_prompt=True, enable_thinking=False
        )

        if "phi" in model_id.lower():
            if prompt_for_model.endswith("<|assistant|>\n"):
                prompt_for_model += "Reasoning:"
            elif prompt_for_model.endswith("<|assistant|>"): 
                 prompt_for_model += "\nReasoning:"
        elif "gemma" in model_id.lower() and prompt_for_model.endswith("<start_of_turn>model\n"):
            prompt_for_model += "Reasoning:"
        elif "llama-3" in model_id.lower() and prompt_for_model.endswith("<|start_header_id|>assistant<|end_header_id|>\n\n"):
             prompt_for_model += "Reasoning:"
        elif "qwen" in model_id.lower() and prompt_for_model.endswith("<|im_start|>assistant\n"):
            prompt_for_model += "Reasoning:"


    except Exception as e:
        print(f"Warning: tokenizer.apply_chat_template failed: {e}. Using manual concatenation (model-specific fallback needed).")
        prompt_for_model = ""
        system_content = ""
        user_content = ""

        for msg in messages:
            if msg['role'] == 'system':
                system_content = msg['content']
            elif msg['role'] == 'user':
                user_content = msg['content']

        combined_user_content = user_content
        if system_content:
            combined_user_content = f"SYSTEM INSTRUCTIONS:\n{system_content}\n\nUSER TASK:\n{user_content}"

        if "gemma" in model_id.lower():
            # Gemma 3 style: <bos><start_of_turn>user\n{message}<end_of_turn>\n<start_of_turn>model\nReasoning:
            prompt_for_model = tokenizer.bos_token if tokenizer.bos_token else ""
            prompt_for_model += f"<start_of_turn>user\n{combined_user_content}<end_of_turn>\n"
            prompt_for_model += f"<start_of_turn>model\nReasoning:" # 어시스턴트의 응답 시작과 Reasoning:
        elif "qwen" in model_id.lower():
            # Qwen style: <|im_start|>system\n{system}<|im_end|>\n<|im_start|>user\n{user}<|im_end|>\n<|im_start|>assistant\nReasoning:
            prompt_for_model = f"<|im_start|>system\n{system_content}<|im_end|>\n" if system_content else ""
            prompt_for_model += f"<|im_start|>user\n{user_content}<|im_end|>\n<|im_start|>assistant\nReasoning:"
        elif "llama-3" in model_id.lower():
            # Llama 3 style: <bos><|start_header_id|>system<|end_header_id|>\n\n{system}<|eot_id|><|start_header_id|>user<|end_header_id|>\n\n{user}<|eot_id|><|start_header_id|>assistant<|end_header_id|>\n\nReasoning:
            prompt_for_model = tokenizer.bos_token if tokenizer.bos_token else ""
            for msg in messages:
                start_header = getattr(tokenizer, 'special_tokens_map', {}).get('<|start_header_id|>', '<|start_header_id|>')
                end_header = getattr(tokenizer, 'special_tokens_map', {}).get('<|end_header_id|>', '<|end_header_id|>')
                eot = getattr(tokenizer, 'special_tokens_map', {}).get('<|eot_id|>', '<|eot_id|>')

                prompt_for_model += f"{start_header}{msg['role']}{end_header}\n\n{msg['content']}{eot}"

            prompt_for_model += f"{start_header}assistant{end_header}\n\nReasoning:"
        elif "phi" in model_id.lower():
            prompt_for_model = f"<|user|>\n{combined_user_content}<|end|>\n<|assistant|>\nReasoning:"
        else:
            prompt_for_model = f"System: {system_content}\n" if system_content else ""
            prompt_for_model += f"User: {user_content}\nAssistant:\nReasoning:"

    model_max_len = tokenizer.model_max_length or 8192
    max_new_tokens = 500
    context_len_budget = max(64, model_max_len - max_new_tokens - 10)

    inputs = tokenizer(prompt_for_model, return_tensors="pt", padding="longest", truncation=True, max_length=context_len_budget)    
    
    eos_ids_list = []
    if tokenizer.eos_token_id is not None:
        if isinstance(tokenizer.eos_token_id, list):
            eos_ids_list.extend(token_id for token_id in tokenizer.eos_token_id if token_id is not None)
        elif isinstance(tokenizer.eos_token_id, int):
             eos_ids_list.append(tokenizer.eos_token_id)
    
    special_eos_tokens_map = { "qwen": "<|im_end|>", "llama-3": "<|eot_id|>", "phi": "<|end|>" }
    for model_key, token_str in special_eos_tokens_map.items():
        if model_key in model_id.lower():
            try:
                token_id = tokenizer.convert_tokens_to_ids(token_str)
                if token_id is not None and token_id not in eos_ids_list:
                    eos_ids_list.append(token_id)
            except Exception as e:
                print(f"DEBUG: Could not convert special token '{token_str}' for {model_key}: {e}")

    if not eos_ids_list and tokenizer.pad_token_id is not None:
        print(f"[WARN] No EOS token ID(s) found. Using pad_token_id ({tokenizer.pad_token_id}) as EOS.")
        eos_ids_list = [tokenizer.pad_token_id]
    
    if not eos_ids_list:
        config_eos_id = getattr(model.config, "eos_token_id", None)
        if isinstance(config_eos_id, int): eos_ids_list.append(config_eos_id)
        elif isinstance(config_eos_id, list): eos_ids_list.extend(config_eos_id)
        if eos_ids_list: print(f"[INFO] Using EOS token ID(s) {eos_ids_list} from model.config.")
        else: raise ValueError(f"EOS token ID is not set for tokenizer, no pad_token, and not found in model.config for model {model_id}.")
   
    output_sequences = model.generate(
        inputs["input_ids"].to(model.device),
        attention_mask=inputs["attention_mask"].to(model.device),
        do_sample=False, # deterministic
        max_new_tokens=max_new_tokens,
        pad_token_id=tokenizer.pad_token_id if tokenizer.pad_token_id is not None else eos_ids_list[0],
        eos_token_id=eos_ids_list
    )

    def strip_bf_number(text):
        match = re.search(r'(\d.*)', text)
        if match:        
            return match.group(1)
        else:        
            return text

    input_token_len = inputs["input_ids"].shape[1]
    generated_tokens = output_sequences[0][input_token_len:]
    full_model_response = tokenizer.decode(generated_tokens, skip_special_tokens=True).strip()
    
    if "phi" in model_id.lower() and full_model_response.startswith("\n"):
        full_model_response = full_model_response.lstrip('\n')
    
    print(f"LLM Full Raw Output for ID {utterance_id_for_output}:\n---\n{full_model_response}\n---")

    current_id_str = str(utterance_id_for_output).strip()
    default_error_line = f"{current_id_str} : {dialogue_info.get('current_speaker','ERR_SPEAKER')} : PARSE_ERROR : NA"
    final_output_line = default_error_line
    reasoning_text = ""

    if "FINAL_CLASSIFICATION:" in full_model_response:        
        reasoning_part, classification_content_after_marker = full_model_response.split("FINAL_CLASSIFICATION:", 1)
                
        reasoning_text = reasoning_part.strip()
        if reasoning_text.lower().startswith("reasoning:"):
            reasoning_text = reasoning_text[len("reasoning:"):].strip()

        potential_lines = strip_bf_number(classification_content_after_marker).strip().splitlines()
        
        found_classification_line_raw = ""
        for line_candidate in potential_lines:
            line_candidate_stripped = line_candidate.strip()
            if line_candidate_stripped.startswith(current_id_str) and line_candidate_stripped.count(':') == 4:
                found_classification_line_raw = line_candidate_stripped
                break 
            elif not found_classification_line_raw and line_candidate_stripped: 
                found_classification_line_raw = line_candidate_stripped


        if found_classification_line_raw:
            cleaned_line = found_classification_line_raw 
            
            match = re.search(rf"({re.escape(current_id_str)}\s*:.*)", cleaned_line)
            if match:
                cleaned_line = match.group(1)
            else: 
                cleaned_line = re.sub(r"^[*\s#\-:!]+", "", cleaned_line)

            expected_prefix = f"{current_id_str} :"
            num_colons = cleaned_line.count(':')
            
            if cleaned_line.startswith(expected_prefix) and num_colons == 3:
                final_output_line = cleaned_line
            else:
                warning_msg = f"  [WARN] Parsed line format mismatch for ID {current_id_str}. "
                if not cleaned_line.startswith(expected_prefix):
                    actual_start_cleaned = cleaned_line.split(':', 1)[0].strip() if ':' in cleaned_line else cleaned_line
                    warning_msg += f"Expected cleaned line to start with ID '{current_id_str}' (prefix: '{expected_prefix}'), but actual start is '{actual_start_cleaned}'. "
                if num_colons != 3:
                    warning_msg += f"Expected 3 colons in cleaned line, got {num_colons}. "
                warning_msg += f"Original Found Raw Line from LLM: '{found_classification_line_raw}'. Attempted Cleaned Line: '{cleaned_line}'"
                print(warning_msg)
        else:
            print(f"  [WARN] 'FINAL_CLASSIFICATION:' found, but no valid classification line after it for ID {current_id_str}. Content after marker: '{classification_content_after_marker.strip()}'")
    else: 
        print(f"  [WARN] 'FINAL_CLASSIFICATION:' marker NOT found for ID {current_id_str}. Attempting fallback search for a line starting with the ID.")
        lines_in_response = full_model_response.splitlines()
        found_in_fallback = False
        for line_raw_fallback in reversed(lines_in_response):
            line_raw_fallback = line_raw_fallback.strip()
            if not line_raw_fallback: continue

            cleaned_line_fallback = line_raw_fallback
            match_fallback = re.search(rf"({re.escape(current_id_str)}\s*:.*)", cleaned_line_fallback)
            if match_fallback:
                cleaned_line_fallback = match_fallback.group(1)
            else:
                cleaned_line_fallback = re.sub(r"^[*\s#\-:!]+", "", cleaned_line_fallback)

            expected_prefix_fallback = f"{current_id_str} :"
            num_colons_fallback = cleaned_line_fallback.count(':')

            if cleaned_line_fallback.startswith(expected_prefix_fallback) and num_colons_fallback == 4:
                final_output_line = cleaned_line_fallback
                print(f"  [INFO] Found classification via fallback for ID {current_id_str}: '{final_output_line}' (Original raw line from LLM: '{line_raw_fallback}')")
                found_in_fallback = True
                break
        
        if not found_in_fallback:
             print(f"  [ERROR] Could not find a valid classification line via any method for ID {current_id_str}.")
             final_output_line = f"{current_id_str} : {dialogue_info.get('current_speaker','ERR_SPEAKER')} : NO_VALID_OUTPUT_FOUND : NA"
             reasoning_text = full_model_response.strip()

    return {"classification_line": final_output_line, "reasoning": reasoning_text}

# --- 5. main function ---
def process_dialogue_dataframe(df, model, tokenizer, model_id="unknown"):
    results = []
    
    if 'timestamp' in df.columns and not pd.api.types.is_string_dtype(df['timestamp']):
        try:
            df['timestamp_str'] = df['timestamp'].astype(str) 
        except Exception as e:
            print(f"Warning: Could not convert 'timestamp' column to string: {e}. Assuming it's already string-like or will be handled by parser.")
            df['timestamp_str'] = df['timestamp'] 
    elif 'timestamp' in df.columns:
         df['timestamp_str'] = df['timestamp']
    else:
        print("Warning: 'timestamp' column not found. Timestamps will be 'N/A'.")
        df['timestamp_str'] = "N/A"

    def timestamp_to_seconds(ts_str):
        if pd.isna(ts_str) or ts_str == "N/A" or not isinstance(ts_str, str): return -1
        parts = str(ts_str).split(':')
        try:
            if len(parts) == 2: return int(parts[0])*60 + int(parts[1])
            if len(parts) == 3: return int(parts[0])*3600 + int(parts[1])*60 + int(parts[2])
        except ValueError:
            return -2 
        return 0 
    
    df['sortable_time'] = df['timestamp_str'].apply(timestamp_to_seconds)
    df['utterance_id'] = df.index.astype(str)

    current_block_index = 0
    last_timestamp_str, last_group, last_date_str = None, None, None
    previous_utterances_in_block = []
    
    for index, row in tqdm(df.iterrows(), total=len(df), desc="Processing Dialogues"):
        current_utterance_id = str(row['utterance_id']) 
        current_group = str(row.get('group', "UnknownGroup"))
        current_date_str = str(row.get('date', "UnknownDate")) 
        current_timestamp_str = str(row.get('timestamp_str', "N/A"))
        
        current_speaker = str(row.get('name', "UnknownSpeaker"))
        current_utterance_text = str(row.get('utterance', ""))

        if (last_group != current_group or 
            last_date_str != current_date_str or 
            last_timestamp_str != current_timestamp_str):
            current_block_index += 1
            previous_utterances_in_block = []

        dialogue_info = {
            "current_block_index": current_block_index,
            "current_group": current_group,
            "current_date": current_date_str, 
            "current_speaker": current_speaker,
            "current_utterance_text": current_utterance_text,
            "current_utterance_indent": 0
        }
        if len(previous_utterances_in_block) >= 1:
            dialogue_info["previous_utterance_1"] = previous_utterances_in_block[-1]['utterance']
            dialogue_info["previous_speaker_1"] = previous_utterances_in_block[-1]['speaker']
        if len(previous_utterances_in_block) >= 2:
            dialogue_info["previous_utterance_2"] = previous_utterances_in_block[-2]['utterance']
            dialogue_info["previous_speaker_2"] = previous_utterances_in_block[-2]['speaker']
        if len(previous_utterances_in_block) >= 3: # 3번째 이전 발화 추가
            dialogue_info["previous_utterance_3"] = previous_utterances_in_block[-3]['utterance']
            dialogue_info["previous_speaker_3"] = previous_utterances_in_block[-3]['speaker']

        classification_result_line = f"{current_utterance_id} : {current_speaker} : CLASSIFICATION_FAILED : NA"
        reasoning_text = ""
        
        if not current_utterance_text.strip():
            classification_result_line = f"{current_utterance_id} : {current_speaker} : Uncorrelated : NA"
            reasoning_text = ""
        else:
            try:
                result = generate_text_classification(
                    model, tokenizer, dialogue_info, current_utterance_id, model_id=model_id
                )
                classification_result_line = result["classification_line"]
                reasoning_text = result["reasoning"]
            except Exception as e:
                print(f"  [CRITICAL ERROR] Classification failed for utterance ID {current_utterance_id}: '{current_utterance_text[:50]}...' - {e}")
                import traceback; traceback.print_exc()
                reasoning_text = ""
        
        results.append({
            'utterance_id': current_utterance_id,
            'timestamp': current_timestamp_str,
            'name': current_speaker,
            'utterance': current_utterance_text,
            'group': current_group,
            'date': current_date_str,
            'cls_original': row.get('cls', 'N/A'),
            'block_index': current_block_index,
            'llm_classification_output_line': classification_result_line,
            'llm_reasoning': reasoning_text
        })

        previous_utterances_in_block.append({'speaker': current_speaker, 'utterance': current_utterance_text})
        if len(previous_utterances_in_block) > 3:
            previous_utterances_in_block.pop(0)
        
        last_timestamp_str, last_group, last_date_str = current_timestamp_str, current_group, current_date_str
            
        if (index + 1) % 20 == 0: 
            if torch.cuda.is_available():
                torch.cuda.empty_cache()
    
    if torch.cuda.is_available(): 
        torch.cuda.empty_cache()

    return pd.DataFrame(results)

# --- 6. LLM output parsing ---
def parse_llm_output_line(line_str, expected_id_str, reasoning=None):
    parts = line_str.split(' : ')
    if len(parts) == 4:
        parsed_id = parts[0].strip()
        parsed_name = parts[1].strip()
        llm_label_full = parts[2].strip() 
        llm_reference = parts[3].strip()
        llm_label_main = llm_label_full
        llm_engagement = "NA" 
        if "Engage Others Idea" in llm_label_full:
            if "(LOW)" in llm_label_full:
                llm_label_main = "Engage Others Idea"
                llm_engagement = "LOW"
            elif "(MEDIUM)" in llm_label_full:
                llm_label_main = "Engage Others Idea"
                llm_engagement = "MEDIUM"
            elif "(HIGH)" in llm_label_full:
                llm_label_main = "Engage Others Idea"
                llm_engagement = "HIGH"
        result_dict = {
            "parsed_utterance_id": parsed_id,
            "parsed_name": parsed_name,
            "parsed_indent": "0",  # Default value since it's no longer in the input
            "llm_label": llm_label_main,      
            "llm_engagement": llm_engagement, 
            "llm_reference": llm_reference,
            "llm_reasoning": reasoning if reasoning is not None else ""
        }
        if parsed_id == expected_id_str:
            return result_dict
        else:
            print(f"  [PARSE_ID_MISMATCH] Expected ID {expected_id_str} but got {parsed_id} in line: '{line_str}'")
            result_dict["llm_label"] = f"ID_MISMATCH_{llm_label_main}"
            return result_dict
    print(f"  [PARSE_WARN] LLM output line does not have 4 parts. Line: '{line_str}' for ID {expected_id_str} (Found {len(parts)} parts)")
    return {
        "parsed_utterance_id": str(expected_id_str), 
        "parsed_name": "ERR_FORMAT",
        "parsed_indent": "0",  # Default value
        "llm_label": "LINE_FORMAT_ERROR", 
        "llm_engagement": "NA",
        "llm_reference": "NA",
        "llm_reasoning": reasoning if reasoning is not None else ""
    }

if __name__ == '__main__':     
    parser = argparse.ArgumentParser(description="Process dialogue transcripts for utterance classification.")
    parser.add_argument(
        "--model_id", type=str, default="meta-llama/Meta-Llama-3.1-8B-Instruct", 
        help="Hugging Face model ID (e.g., 'meta-llama/Meta-Llama-3.1-8B-Instruct')."
    )
    parser.add_argument("--transcript_dir", type=str, default="./transcripts/", help="Directory with .docx transcripts.")
    parser.add_argument("--output_csv", type=str, default="classification_results.csv", help="Output CSV filename.")

    args = parser.parse_args()   

    start_time = time.time()
    print(f"Using model: {args.model_id}")
    print(f"Transcript directory: {args.transcript_dir}")
    print(f"Output CSV: {args.output_csv}")

    try:
        model, tokenizer = load_model_and_tokenizer(args.model_id)   
        print(model.hf_device_map)
        opath = args.transcript_dir
        os.makedirs(opath, exist_ok=True)
        if not os.listdir(opath) and opath == "./test_transcripts/":
            print(f"Creating dummy docx files in {opath} for testing...")
            dummy_doc_content1="""0:01\nSpeakerA: Hello.\n0:02\nSpeakerB: Hi?\n0:03\nSpeakerA: Add 2 and 2.\n0:05\nSpeakerC: 4.\n0:06\nSpeakerA: Yes."""
            doc1=docx.Document(); [doc1.add_paragraph(l.strip()) for l in dummy_doc_content1.strip().split('\n')]; doc1.save(os.path.join(opath,"AL_G1_Test_1.1_f.docx"))
            print("Dummy file created.")

        fpath = sorted([f for f in os.listdir(opath) if f.endswith(".docx") and not f.startswith("~")])
        if not fpath: print(f"No .docx files found in {opath}. Exiting."); exit()
        print(f"Found {len(fpath)} docx files.")

        group_list = ["total" if ("ipad"in f.lower())or("finalized"in f.lower()) else f.split("_")[1] if len(f.split("_"))>1 else "grp_err" for f in fpath]
        date_list = [re.findall(r'\d{1,2}\.\d{1,2}',f)[0] if re.findall(r'\d{1,2}\.\d{1,2}',f) else "date_err" for f in fpath]
        class_list = [f[0:2] if len(f)>=2 else "cls_err" for f in fpath]
        
        full_path = pd.DataFrame({"group":group_list,"date":date_list,"cls":class_list,"path":[os.path.join(opath,f) for f in fpath]})
        full_path['group'] = full_path['group'].str.replace(")", "", regex=False).str.replace("(", "_", regex=False)
        
        full_path_filtered = full_path[(full_path['cls'] == "AL") & (full_path['group'] != "total")].reset_index(drop=True)
        N = len(full_path_filtered)
        if N == 0: print("No files after filtering. Exiting."); exit()
        print(f"Processing {N} files after filtering.")

        data_list = []
        for _, row_fp in tqdm(full_path_filtered.iterrows(), total=N, desc="Extracting Docx"):
            try: data_list.append(extract_dialogue_from_docx(row_fp['path']).assign(group=row_fp['group'],date=row_fp['date'],cls=row_fp['cls']))
            except Exception as e: print(f"Error extracting {row_fp['path']}: {e}")
        
        if not data_list: print("No data extracted. Exiting."); exit()
        df_input = pd.concat(data_list, ignore_index=True)
        if df_input.empty: print("Concat DataFrame empty. Exiting."); exit()
        print(f"\nTotal utterances to process: {len(df_input)}")
        
        df_results = process_dialogue_dataframe(df_input.copy(), model, tokenizer, model_id=args.model_id)

        print("\n--- Processing Complete ---")

        parsed_data = [ {**row_res.to_dict(), **parse_llm_output_line(str(row_res['llm_classification_output_line']), str(row_res['utterance_id']), row_res['llm_reasoning'])} for _, row_res in df_results.iterrows()]
        df_final_results = pd.DataFrame(parsed_data)
        df_final_results.drop(columns=['utterance', 'reason'], errors='ignore', inplace=True)
        df_final_results.dropna(how='all', axis=1, inplace=True)
        df_final_results.to_csv(args.output_csv, index=False, encoding='utf-8-sig') 
        print(f"\nResults saved to {args.output_csv}")

        total_time = time.time() - start_time
        print(f"Total execution time: {total_time:.2f}s ({total_time/60:.2f}min).")
        if len(df_input)>0: print(f"Avg time/utterance: {total_time/len(df_input):.2f}s.")

    except Exception as e:
        print(f"An critical error occurred in main execution: {e}")
        import traceback; traceback.print_exc()